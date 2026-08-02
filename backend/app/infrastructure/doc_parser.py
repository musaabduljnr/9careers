import io
import logging
from backend.app.domain.interfaces import DocumentParser

logger = logging.getLogger(__name__)

class DocumentParserImpl(DocumentParser):
    def parse_pdf(self, file_bytes: bytes) -> str:
        """Parse text content from a PDF document using pypdf."""
        try:
            from pypdf import PdfReader
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts).strip()
        except Exception as e:
            logger.error(f"Failed to parse PDF file: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def parse_docx(self, file_bytes: bytes) -> str:
        """Parse text content from a DOCX document using python-docx."""
        try:
            import docx
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            # Extract text from tables too
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                        
            return "\n".join(text_parts).strip()
        except Exception as e:
            logger.error(f"Failed to parse DOCX file: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
